"""Build the IEEE R10-HTC conference paper as a Word .docx in the IEEE two-column layout.

Mirrors docs/paper/ieee-r10htc.md (and the IEEEtran .tex) into a Word file styled to the
IEEE conference template: US-Letter, 0.625 in side margins, two-column body with a single
1-column title/author block, Times New Roman, IEEE heading/caption/table/reference sizes,
full-width tables (Table I-III) placed in 1-column section islands, real figures where
committed and a bordered placeholder for the author-drawn bias-bar figure (Fig. 3).

Word is Unicode-native, so symbols (x, >=, ~, arrows, superscripts) are kept verbatim.
Inline **bold** / *italic* markdown in the source strings is honoured by a tiny parser.

Run:  python scripts/build_ieee_docx.py
Out:  docs/paper/ieee-r10htc.docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "docs" / "paper" / "figures"
OUT = ROOT / "docs" / "paper" / "ieee-r10htc.docx"

BODY_PT = 10        # IEEE body text
ABS_PT = 9          # abstract / index terms
TBL_PT = 8          # table body
CAP_PT = 8          # caption / table title
REF_PT = 8          # references
FONT = "Times New Roman"


# ---------------------------------------------------------------- low-level helpers
def _set_run_font(run, size_pt, *, bold=False, italic=False, small_caps=False):
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.small_caps = small_caps
    # force the font for all script ranges (Word ignores .name for some ranges otherwise)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT)


_INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def _add_runs(paragraph, text, size_pt):
    """Add runs to a paragraph, honouring **bold** and *italic* markdown spans."""
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            _set_run_font(paragraph.add_run(tok[2:-2]), size_pt, bold=True)
        elif tok.startswith("*") and tok.endswith("*"):
            _set_run_font(paragraph.add_run(tok[1:-1]), size_pt, italic=True)
        else:
            _set_run_font(paragraph.add_run(tok), size_pt)


def body_para(doc, text, *, size_pt=BODY_PT, indent=True, justify=True, space_after=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    if indent:
        pf.first_line_indent = Inches(0.2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    _add_runs(p, text, size_pt)
    return p


def centered(doc, text, size_pt, *, bold=False, italic=False, small_caps=False, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    _set_run_font(p.add_run(text), size_pt, bold=bold, italic=italic, small_caps=small_caps)
    return p


def heading(doc, roman, title):
    """IEEE primary heading: centered, small-caps, '<Roman>.  <Title>'."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    _set_run_font(p.add_run(f"{roman}.  {title}"), BODY_PT, small_caps=True)
    return p


def lead_in(doc, bold_lead, rest):
    """Paragraph that opens with a bold run-in phrase (IEEE 'paragraph heading' style)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(0)
    pf.first_line_indent = Inches(0.2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_run_font(p.add_run(bold_lead), BODY_PT, bold=True)
    _add_runs(p, rest, BODY_PT)
    return p


def set_columns(section, num, space_in=0.2):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(int(space_in * 1440)))
    cols.set(qn("w:equalWidth"), "1")


def new_section(doc, num_cols):
    sec = doc.add_section(WD_SECTION.CONTINUOUS)
    set_columns(sec, num_cols)
    return sec


def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def make_table(doc, header, rows, *, col_aligns=None):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    aligns = col_aligns or ["left"] * len(header)
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        _shade(cell, "F2F2F2")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(para.add_run(h), TBL_PT, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            para = cells[j].paragraphs[0]
            para.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }[aligns[j]]
            _add_runs(para, val, TBL_PT)
    return t


def table_caption(doc, label, title):
    centered(doc, label, CAP_PT, small_caps=True, space_before=6, space_after=0)
    centered(doc, title, CAP_PT, space_after=4)


def table_note(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    _add_runs(p, text, 7)


def figure(doc, img_path: Path, caption_label, caption_text, width_in=3.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    if img_path.exists():
        p.add_run().add_picture(str(img_path), width=Inches(width_in))
    else:
        _set_run_font(p.add_run(f"[missing image: {img_path.name}]"), CAP_PT, italic=True)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.line_spacing = 1.0
    _set_run_font(cap.add_run(caption_label + " "), CAP_PT, bold=True)
    _add_runs(cap, caption_text, CAP_PT)


def figure_placeholder(doc, caption_label, caption_text, box_text, height_in=1.5):
    box = doc.add_table(rows=1, cols=1)
    box.style = "Table Grid"
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.rows[0].cells[0]
    cell.width = Inches(3.4)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(para.add_run(box_text), CAP_PT, italic=True)
    tr = box.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(height_in * 1440)))
    h.set(qn("w:hRule"), "atLeast")
    trPr.append(h)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_after = Pt(6)
    _set_run_font(cap.add_run(caption_label + " "), CAP_PT, bold=True)
    _add_runs(cap, caption_text, CAP_PT)


# ---------------------------------------------------------------- document
def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_PT)

    s0 = doc.sections[0]
    s0.page_width, s0.page_height = Inches(8.5), Inches(11)
    s0.top_margin = Inches(0.75)
    s0.bottom_margin = Inches(1.0)
    s0.left_margin = Inches(0.625)
    s0.right_margin = Inches(0.625)
    set_columns(s0, 1)

    # ---- title / author block (1 column) ----
    centered(doc, "Open, Reproducible 30 m Multi-Hazard Flood Screening "
                  "for Under-Resourced Southeast Asian Cities",
             20, space_before=0, space_after=10)
    centered(doc, "[Author Name(s) — TBD]", 11, space_after=0)
    centered(doc, "[Affiliation — TBD]", 10, italic=True, space_after=0)
    centered(doc, "[Department, Institution, City, Country]", 10, italic=True, space_after=0)
    centered(doc, "[email — TBD]", 10, italic=True, space_after=8)

    # ---- switch to two columns ----
    new_section(doc, 2)

    # Abstract
    ab = doc.add_paragraph()
    ab.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ab.paragraph_format.line_spacing = 1.0
    ab.paragraph_format.first_line_indent = Inches(0.2)
    _set_run_font(ab.add_run("Abstract—"), ABS_PT, bold=True, italic=True)
    _set_run_font(ab.add_run(ABSTRACT), ABS_PT, italic=True)

    ix = doc.add_paragraph()
    ix.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ix.paragraph_format.first_line_indent = Inches(0.2)
    ix.paragraph_format.space_before = Pt(6)
    _set_run_font(ix.add_run("Index Terms—"), ABS_PT, bold=True, italic=True)
    _set_run_font(ix.add_run(INDEX_TERMS), ABS_PT, italic=True)

    # I. Introduction
    heading(doc, "I", "Introduction")
    body_para(doc, INTRO_1)
    body_para(doc, INTRO_2)
    body_para(doc, INTRO_3)

    # Table I (full width island)
    new_section(doc, 1)
    table_caption(doc, "TABLE I", "Comparator Landscape for ASEAN-Relevant Flood-Screening Tools")
    make_table(doc, TBL1_HEAD, TBL1_ROWS,
               col_aligns=["left", "left", "center", "center", "center", "center"])
    table_note(doc, TBL1_NOTE)
    new_section(doc, 2)

    # II. Open Multi-Hazard Pipeline
    heading(doc, "II", "Open Multi-Hazard Pipeline")
    lead_in(doc, "Open-data inputs. ", PIPE_INPUTS)
    lead_in(doc, "Per-country IDF anchoring. ", PIPE_IDF)
    lead_in(doc, "Three hazards. ", PIPE_HAZARDS)
    lead_in(doc, "Scenarios. ", PIPE_SCEN)
    figure(doc, FIG / "ieee_fig1_pipeline.png", "Fig. 1.", FIG1_CAP, width_in=2.85)

    # III. The Open Flood Atlas
    heading(doc, "III", "The Open Flood Atlas")
    body_para(doc, ATLAS_1)
    body_para(doc, ATLAS_2)

    new_section(doc, 1)
    table_caption(doc, "TABLE II", "Combined Flood Extent at RP100, SSP5-8.5 / 2100 (km²)")
    make_table(doc, TBL2_HEAD, TBL2_ROWS,
               col_aligns=["left", "right", "right", "right", "right", "left"])
    table_note(doc, TBL2_NOTE)
    new_section(doc, 2)

    figure(doc, FIG / "ieee_fig2_bangkok_rp100.png", "Fig. 2.", FIG2_CAP)

    # IV. Validation and Trustworthiness
    heading(doc, "IV", "Validation and Trustworthiness")
    body_para(doc, VALID_1)

    new_section(doc, 1)
    table_caption(doc, "TABLE III", "Documented-Hotspot Gate (Present-Day, RP100, ≥ 0.10 m, 50 m Radius)")
    make_table(doc, TBL3_HEAD, TBL3_ROWS,
               col_aligns=["left", "center", "center", "center", "center", "left"])
    table_note(doc, TBL3_NOTE_CC)
    table_note(doc, TBL3_NOTE_D)
    new_section(doc, 2)

    body_para(doc, VALID_2)
    body_para(doc, VALID_3)
    figure(doc, FIG / "ieee_fig3_bathtub_bias.png", "Fig. 3.", FIG3_CAP)

    # V. Impact and Limitations
    heading(doc, "V", "Impact and Limitations")
    body_para(doc, IMPACT_1)
    body_para(doc, IMPACT_2)

    # VI. Conclusion
    heading(doc, "VI", "Conclusion")
    body_para(doc, CONCLUSION)

    # Reproducibility
    heading(doc, "", "Reproducibility")
    body_para(doc, REPRO)

    # References
    heading(doc, "", "References")
    for i, ref in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        _set_run_font(p.add_run(f"[{i}] "), REF_PT)
        _add_runs(p, ref, REF_PT)

    doc.save(OUT)
    # round-trip sanity: re-open to confirm a structurally valid docx
    Document(str(OUT))
    print(f"wrote {OUT}  ({OUT.stat().st_size // 1024} KB)")


# ---------------------------------------------------------------- content (verbatim from ieee-r10htc.md)
ABSTRACT = (
    "The Southeast Asian cities most exposed to flooding are precisely those least able to obtain "
    "usable hazard information. The only comparable-resolution (30 m) three-hazard model with regional "
    "coverage is commercial and closed, and its free-access tranche explicitly excludes the major ASEAN "
    "megacities; bespoke per-city engineering studies are not publicly reproducible; and the open global "
    "products are an order of magnitude coarser and omit the pluvial (rain-driven) hazard that dominates "
    "urban flooding in the region. We present an open-source, open-data pipeline that produces "
    "design-event coastal, fluvial, and pluvial flood-depth maps at 30 m for four cities — Singapore, "
    "Kuala Lumpur, Bangkok, and Jakarta — across four ASEAN countries, under four climate "
    "combinations (SSP2-4.5 and SSP5-8.5 × 2050 and 2100). Every input is freely accessible without "
    "registration, and the pluvial hazard is anchored to the national meteorological services' published "
    "Intensity–Duration–Frequency (IDF) standards, eliminating the 28–62 % deficit that global synthetic "
    "rainfall carries over tropical-convective extremes. We make the open model trustworthy, not merely "
    "reproducible: a model-blind documented-hotspot location-skill gate yields statistically significant "
    "discriminative skill in all four cities tested, and a bathtub-bias characterisation (coastal "
    "over-prediction of 1.7–25× at RP100) is closed by a local-inertia shallow-water solver (Bangkok "
    "over-prediction 12.5× → ≈1×). The release is intended as a public good for under-resourced municipal "
    "disaster-management and climate-adaptation agencies."
)

INDEX_TERMS = ("humanitarian technology, flood risk, open data, multi-hazard, decision support, "
               "climate adaptation, Southeast Asia.")

INTRO_1 = (
    "Southeast Asia contains six of the ten cities globally most exposed to coastal flooding by 2100 "
    "under high-emission scenarios. Roughly 750 million people live in the ASEAN region, with about a "
    "fifth of regional GDP on flood-exposed land [1], and the stressors compound: AR6 P50 sea-level "
    "rise under SSP5-8.5 by 2100 ranges from +0.62 m (Singapore) to +1.62 m (Bangkok); tropical "
    "convective rainfall is intensifying at or above the Clausius–Clapeyron rate [2]; and several "
    "megacities subside at 1–25 cm yr⁻¹ from groundwater extraction and clay compaction [3]–[5]. The "
    "2011 Thailand megaflood, the 2020 Jakarta monsoon floods, and the 2021 Kuala Lumpur flash floods "
    "are recent reminders that the exposure is multi-hazard — coastal, fluvial, and pluvial — and "
    "concurrent [6], [7]."
)
INTRO_2 = (
    "Yet the cities carrying this risk cannot readily obtain affordable, usable hazard maps. Existing "
    "flood-screening tools fall into three categories, none of which simultaneously offers high "
    "resolution, multi-hazard scope, open code, open data, and per-country calibration (Table I). The "
    "only 30 m three-hazard global model is commercial [8], and its free-access tranche is restricted "
    "to a set of developing countries that *excludes the major ASEAN megacities*. Bespoke municipal "
    "engineering studies exist but are closed and not reproducible. The open global alternatives are an "
    "order of magnitude coarser (~10 km) and lack the pluvial layer that drives most urban flood damage "
    "in the region [9]–[11]. The result is an equity gap: open, high-resolution, multi-hazard flood "
    "information — a basic input to humanitarian disaster-risk reduction and adaptation planning — is "
    "unavailable to exactly the agencies that most need it."
)
INTRO_3 = (
    "This paper addresses that gap. We present an open-source, open-data, per-country-calibrated 30 m "
    "multi-hazard flood pipeline and atlas for ASEAN cities, and — crucially — we validate that it "
    "locates real floods. The paper covers four cities (Singapore, Kuala Lumpur, Bangkok, Jakarta) "
    "across four countries; two further cities (Manila, Ho Chi Minh City) are deferred pending a solver "
    "extension (Section V)."
)

TBL1_HEAD = ["Tool", "Resolution", "Hazards", "Open code", "Open data", "Per-country IDF"]
TBL1_ROWS = [
    ["**This work**", "**30 m**", "**C + F + P**", "**Yes**", "**Yes**", "**Yes (4 services)**"],
    ["Fathom 3.0 [8]", "30 m", "C + F + P", "No", "No", "Global synthetic †"],
    ["Aqueduct Floods 4.0 [9],[10]", "~10 km", "C + F", "Method only", "Yes", "None"],
    ["GLOFRIS / PCR-GLOBWB [11]", "~10 km", "F (+C)", "Yes", "Yes", "None"],
    ["City engineering studies", "High", "Various", "No", "No", "Per-city"],
]
TBL1_NOTE = ("† Fathom's free-access tranche is restricted to selected developing countries and excludes "
             "the major ASEAN megacities covered here — the equity gap this work targets. "
             "C = coastal, F = fluvial, P = pluvial.")

PIPE_INPUTS = (
    "A per-city configuration drives a single orchestration script through the pipeline. Every input is "
    "free and most require no registration: the Copernicus GLO-30 DEM (via the Microsoft Planetary "
    "Computer), ERA5-Land precipitation [12] (Open-Meteo), University of Hawaii Sea Level Center (UHSLC) "
    "tide gauges, GloFAS v4 discharge reanalysis [13] (Open-Meteo Flood API), IPCC AR6 sea-level "
    "projections [14] (NCAR/Rutgers Zarr), ESA WorldCover 2021 land cover [15], and OpenStreetMap "
    "waterways. The implementation is pure Python on the scientific stack (rasterio, scipy, pyproj, "
    "pysheds, numba); no external binaries or licensed data are required, which is what makes the "
    "pipeline reproducible from scratch by a third party."
)
PIPE_IDF = (
    "The single most important in-region calibration is the pluvial design rainfall. Global synthetic "
    "rainfall statistics under-represent tropical-convective extremes by 28–62 % relative to the "
    "national IDF curves, so we fit a two-anchor Gumbel distribution directly to each country's "
    "published design standard: PUB (Singapore), JPS-MSMA (Malaysia), TMD-RID (Thailand), and BMKG "
    "(Indonesia). The storm duration is matched to the dominant local flood mechanism — Singapore uses "
    "a 1-hour secondary-drain IDF (its documented flash floods are sub-hourly convective bursts), the "
    "others a 6-hour IDF — and the post-drain excess rainfall is the depth the drainage network cannot "
    "convey."
)
PIPE_HAZARDS = (
    "*Coastal:* a GEV fit to UHSLC annual-maximum storm-surge residuals (T_TIDE-detided [16]), "
    "consistent with global storm-surge reanalyses [17], [18], plus the AR6 SLR delta, brought onto the "
    "EGM2008 DEM datum via a CMEMS mean-dynamic-topography offset, and routed by a local-inertia "
    "shallow-water solver [19] on a staggered grid; a connectivity-based bathtub solver is retained as a "
    "fallback where the sea is enclosed within the DEM domain. *Fluvial:* GloFAS-derived design stage "
    "with bankfull subtraction for rivers carrying permanent baseflow, mapped via Height Above Nearest "
    "Drainage (HAND) [20], [21] referenced to the *main-stem trunk the modelled discharge represents* — "
    "flow-accumulation channels at the GloFAS sub-basin (catchment) scale rather than a "
    "channel-initiation threshold or raw OpenStreetMap network, which over-broaden the floodplain. "
    "*Pluvial:* the IDF excess is routed by a catchment-routed fill-and-spill cascade [22] (depressions "
    "resolved by a Planchon–Darboux fill [23]) with a per-cell runoff coefficient from WorldCover land "
    "cover, producing a return-period-dependent ponding extent rather than a uniform cap. The three "
    "depth rasters are composited by per-pixel maximum."
)
PIPE_SCEN = (
    "All hazards are produced for SSP2-4.5 and SSP5-8.5 at 2050 and 2100, on a subsidence-corrected DEM "
    "(a zone-based correction for the documented post-2013 subsidence in Jakarta and Bangkok). Per-step "
    "methodological detail is documented and reproducible in the open repository."
)
FIG1_CAP = ("Pipeline data-flow schematic: free open-data inputs → per-hazard solvers (inertial coastal, "
            "main-stem-HAND fluvial, fill-and-spill pluvial) → per-pixel-maximum composite and severity "
            "classification.")

ATLAS_1 = (
    "At the canonical RP100 / SSP5-8.5 / 2100 combination, the atlas shows the expected physical "
    "contrast between cities (Table II): low-lying Bangkok is coastal-dominated, Jakarta "
    "fluvial/pluvial-dominated, and inland Kuala Lumpur pluvial-dominated. The combined extents are "
    "*screening upper bounds* under explicit no-pumping, no-sub-pixel-defence assumptions (Section V); "
    "the value of the atlas is less in any single absolute number than in the consistent, comparable "
    "surface it provides across cities, hazards, and scenarios."
)
ATLAS_2 = (
    "The clearest policy-relevant signal is the *mitigation delta* — the flooded area avoided by meeting "
    "a lower-emissions pathway. The avoided Bangkok coastal RP100 land under SSP2-4.5 versus SSP5-8.5 at "
    "2100 is −133 km². Because the model uses the same solver and inputs for every city in every "
    "scenario, such cross-scenario deltas are robust to the absolute bias and are the cleanest single "
    "number the atlas offers to an adaptation planner."
)
TBL2_HEAD = ["City", "Coastal", "Fluvial", "Pluvial", "Combined", "Dominant"]
TBL2_ROWS = [
    ["Singapore", "67", "112 *", "19", "184", "Canal + coastal"],
    ["Kuala Lumpur (core)", "0", "126", "268", "362", "Pluvial"],
    ["Bangkok (klong)", "3,500", "788", "433", "3,598", "Coastal"],
    ["Jakarta", "159", "389", "169", "609", "Fluvial + pluvial"],
]
TBL2_NOTE = ("* Singapore's “fluvial” layer is PUB primary canal-overflow under long-duration design "
             "rainfall, not natural-river flooding. The Bangkok coastal extent is a bathtub upper bound; "
             "the inertial-corrected value is 283 km² (Section IV). All values are current-pipeline "
             "bathtub RP100 outputs; the coastal extents reproduce the documented benchmarks within ~1 %.")
FIG2_CAP = ("Bangkok present-day-baseline RP100 combined-hazard flood-depth map (SSP5-8.5 forcing, 2020 "
            "horizon; ~1,374 km² above 0.1 m), illustrating the screening flood envelope on the flat "
            "Chao Phraya delta. Depth classes: minor (0.1–0.15 m), moderate (0.15–0.5 m), major "
            "(0.5–1 m), severe (>1 m).")

VALID_1 = (
    "An open model is only useful if it can be trusted to flood where flooding actually occurs. We test "
    "this directly with a *model-blind documented-hotspot location-skill gate*. For each city we freeze, "
    "before consulting any model output, a register of localities with *documented* flooding (positives) "
    "and localities documented to have stayed *dry* (controls), each geocoded and DEM-verified. The "
    "combined present-day wet mask (pluvial ∨ fluvial ∨ coastal, ≥ 0.10 m, within a 50 m hit radius "
    "matching the geocoding precision) is scored against the register at the event-matched RP100, "
    "reporting hit-rate (HR; sensitivity), correct-reject-rate (CRR; specificity), and the "
    "Peirce–Hanssen–Kuipers true skill statistic (TSS = HR + CRR − 1) with a bootstrap 95 % confidence "
    "interval. The register is a *consistency check*: parameters are anchored upstream to documented "
    "facts, and the gate is never used to retune them."
)
TBL3_HEAD = ["City", "pos/dry", "HR", "CRR", "TSS [95 % CI]", "verdict"]
TBL3_ROWS = [
    ["Kuala Lumpur", "17/7", "0.76", "0.86", "**0.62** [0.25, 0.88]", "PASS"],
    ["Singapore", "38/20", "0.82", "0.65", "**0.47** [0.21, 0.72]", "fail-CRR ‡"],
    ["Bangkok", "16/7", "0.56", "0.86", "**0.42** [0.04, 0.75]", "fail-HR"],
    ["Jakarta", "18/8", "0.89", "0.50", "**0.39** [0.03, 0.75]", "fail-CRR †"],
]
TBL3_NOTE_CC = (
    "‡ Singapore is the city in which the gate was first developed; it is scored here on the *identical* "
    "basis as the other three, re-aligned from the earlier pluvial-only / RP50 / 150 m convention — the "
    "model itself is unchanged. Its 58-point register is the largest, its positives are PUB List of "
    "Flood-Prone Areas localities geocoded from road names (medium confidence, hence the conservative "
    "50 m radius), and its CRR shortfall is the combined RP100 wet mask catching low-lying dry controls "
    "the deliberately conservative pluvial-only layer alone would spare."
)
TBL3_NOTE_D = (
    "† Jakarta's TSS *before* the model-blind reclassification of two documented-flooded mislabels "
    "(Menteng, Gambir — both inundated in 2007 and 2013) was 0.16 (no skill); the reclassification is "
    "anchored to the flood record, not the gate."
)
VALID_2 = (
    "All four cities reach *statistically significant discriminative skill* — every TSS confidence "
    "interval excludes zero. The harness originated on Singapore and transferred to Kuala Lumpur, "
    "Bangkok and Jakarta with no engine changes. Two findings generalise. First, a *main-stem-HAND "
    "rule*: HAND must be referenced to the trunk channel the modelled discharge represents (for Kuala "
    "Lumpur, the ≥ 180 km² accumulation trunk), not a channel-initiation or full-OSM network; doing so "
    "fixes a false-positive on a 60–77 m hill by correct physics and yields a credible floodplain, "
    "lifting Kuala Lumpur to CRR 0.86 / TSS 0.62 (PASS). Second, a *dry-control discipline*: a genuine "
    "control that the model floods stays in the register as a reported false positive, never dropped to "
    "pass; only a control that is *independently documented-flooded* is corrected — Jakarta's "
    "central-levee controls (Menteng, Gambir) sit on the Ciliwung corridor and are documented inundated "
    "in 2007 and 2013, so reclassifying them on the flood-record evidence (decided model-blind) lifts "
    "Jakarta from no-skill (TSS 0.16) to significant skill (TSS 0.39). The remaining shortfalls are "
    "*documented structural limits, not tuning failures*: Bangkok's HR is bounded because the 2011 "
    "reference flood was sourced from a catchment whose headwaters lie outside the model domain, and "
    "Jakarta's residual specificity loss is fill-and-spill over-ponding on genuinely elevated ground."
)
VALID_3 = (
    "A second, complementary result establishes that the open model's coastal layer is corrected, not "
    "merely cheap. A bathtub solver — the default in open screening tools — over-predicts documented "
    "present-day coastal inundation by *1.7–25× at RP100*, because 30 m terrain cannot resolve the "
    "sub-pixel road raises, canals, and bunds that protect small present-day events. Replacing it with "
    "the local-inertia shallow-water solver, where the sea connects to the domain boundary, brings the "
    "over-prediction to ≈1×: Bangkok's RP100 coastal extent drops from 3,546 km² to *283 km² (a 12.5× "
    "reduction)*, within ~30 % of the documented 2011 megaflood extent. The bias correction is a "
    "solver-architecture fix, not a data limitation (Fig. 3). Two further checks are consistent and "
    "space-limited here: the IDF-anchor re-derivation is exact (0 fails across all configurations), the "
    "documented high-water-mark depth cross-check is in-band at three of five points, and satellite "
    "extent-CSI is reported only as an observation-limited sanity check (urban SAR layover and MODIS "
    "coarseness suppress the reference)."
)
FIG3_CAP = ("Bathtub-bias factor (model / documented) at RP2 and RP100 by city, with the local-inertia "
            "overlay for the three solver-compatible cities (log scale); the 12.5× Bangkok RP100 "
            "reduction to ≈1× is the headline. Singapore's residual ratio stays high because its "
            "documented present-day coastal extent is near zero.")

IMPACT_1 = (
    "The intended use is humanitarian: open, validated, reproducible hazard maps as decision-support for "
    "the municipal disaster-management and climate-adaptation agencies that cannot license commercial "
    "models. Because every input is free and every parameter is documented, an agency or a regional "
    "university can rebuild the atlas for its own city, interrogate the assumptions, and extend it — "
    "lowering the barrier to first-order flood-risk information from a procurement question to a "
    "software install."
)
IMPACT_2 = (
    "The model is honest as a *screening upper bound* under three assumptions: no active pumping, no "
    "sub-pixel defences resolved by the 30 m DEM, and per-pixel-maximum (marginal, not joint-exceedance) "
    "multi-hazard composition. The scenario forcing is consistent across the full SSP × horizon grid — "
    "monotone in scenario severity and within physical plausibility bounds, verified by an automated "
    "consistency guard — and the headline quantitative results use the validated SSP5-8.5/2100 and "
    "present-day cells. Two further findings bound transferability and are reported rather than tuned "
    "away: single-stage HAND does not transfer to flat deltas fed by an out-of-domain mega-river "
    "(Bangkok, Jakarta) [24], and the pluvial solver is currently heterogeneous across cities (a "
    "documented homogenisation step). The two further ASEAN cities — Manila (enclosed bay) and Ho Chi "
    "Minh City (enclosed delta) — are deferred because their enclosed-sea topology breaks the inertial "
    "solver's wall condition; relaxing that condition is the principal prerequisite for extending the "
    "bias-aware coastal treatment to them."
)
CONCLUSION = (
    "We have shown that an open-source, open-data, per-country-calibrated 30 m multi-hazard flood atlas "
    "for ASEAN cities is feasible, reproducible from free data, and — by a model-blind location-skill "
    "gate — demonstrably skilful in the four cities tested, with its coastal over-prediction "
    "structurally corrected. To our knowledge it is the first such atlas released with full code, "
    "configuration, and validation transparency for the region. By targeting exactly the cities that "
    "commercial high-resolution models exclude, the release is offered as a public good for humanitarian "
    "flood-risk reduction, and as an invitation to extend."
)
REPRO = ("Code, per-city configuration, and the validation registers are released openly [repository "
         "URL — TBD], enabling third-party rebuild of every result from free data.")

# Ordered by first citation in the text (IEEE convention).
REFERENCES = [
    "Asian Development Bank, “Southeast Asia and the economics of global climate stabilization,” 2022.",
    "D. P. Lenderink et al., “Scaling of extreme rainfall with temperature,” 2017.",
    "D. Chaussard, F. Amelung, H. Abidin, and S.-H. Hong, “Sinking cities in Indonesia: ALOS PALSAR "
    "detects rapid subsidence,” Remote Sens. Environ., vol. 128, pp. 150–161, 2013.",
    "H. Z. Abidin et al., “Land subsidence of Jakarta and its relation with urban development,” Nat. "
    "Hazards, vol. 59, pp. 1753–1771, 2011.",
    "N. Phien-wej, P. H. Giao, and P. Nutalaya, “Land subsidence in Bangkok, Thailand,” Eng. Geol., vol. "
    "82, pp. 187–201, 2006.",
    "S. Hallegatte, C. Green, R. J. Nicholls, and J. Corfee-Morlot, “Future flood losses in major "
    "coastal cities,” Nat. Clim. Change, vol. 3, pp. 802–806, 2013.",
    "B. Tellman et al., “Satellite imaging reveals increased proportion of population exposed to "
    "floods,” Nature, vol. 596, pp. 80–86, 2021.",
    "O. E. J. Wing et al., “A 30 m global flood inundation model for all climates,” 2024.",
    "H. Hofste et al., “Aqueduct Floods methodology,” World Resources Institute Tech. Note, 2019.",
    "P. J. Ward et al., “Aqueduct Floods: global flood risk maps and analysis,” 2020.",
    "E. H. Sutanudjaja et al., “PCR-GLOBWB 2: a 5 arcmin global hydrological and water resources "
    "model,” Geosci. Model Dev., vol. 11, pp. 2429–2453, 2018.",
    "J. Muñoz-Sabater et al., “ERA5-Land: a state-of-the-art global reanalysis dataset for land "
    "applications,” Earth Syst. Sci. Data, vol. 13, pp. 4349–4383, 2021.",
    "L. Alfieri et al., “A global network for operational flood risk reduction (GloFAS),” Environ. Sci. "
    "Policy, vol. 84, pp. 149–158, 2018.",
    "B. Fox-Kemper et al., “Ocean, Cryosphere and Sea Level Change,” in Climate Change 2021: The "
    "Physical Science Basis (IPCC AR6 WGI), Cambridge Univ. Press, 2021.",
    "D. Zanaga et al., “ESA WorldCover 10 m 2021 v200,” Zenodo, 2022.",
    "R. Pawlowicz, B. Beardsley, and S. Lentz, “Classical tidal harmonic analysis including error "
    "estimates in MATLAB using T_TIDE,” Comput. Geosci., vol. 28, pp. 929–937, 2002.",
    "S. Muis, M. Verlaan, H. C. Winsemius, J. C. J. H. Aerts, and P. J. Ward, “A global reanalysis of "
    "storm surges and extreme sea levels,” Nat. Commun., vol. 7, 11969, 2016.",
    "S. Muis et al., “A high-resolution global dataset of extreme sea levels, tides, and storm surges, "
    "including future projections,” Front. Mar. Sci., vol. 7, 263, 2020.",
    "P. D. Bates, M. S. Horritt, and T. J. Fewtrell, “A simple inertial formulation of the shallow water "
    "equations for efficient two-dimensional flood inundation modelling,” J. Hydrol., vol. 387, no. 1–2, "
    "pp. 33–45, 2010.",
    "A. D. Nobre et al., “Height Above the Nearest Drainage — a hydrologically relevant new terrain "
    "model,” J. Hydrol., vol. 404, pp. 13–29, 2011.",
    "B. Schwanghart and D. Scherler, “TopoToolbox 2,” Earth Surf. Dynam., vol. 2, pp. 1–7, 2014.",
    "R. Barnes, K. L. Callaghan, and A. D. Wickert, “Computing water flow through complex landscapes — "
    "Part 2: Fill–Spill–Merge,” Earth Surf. Dynam., vol. 8, pp. 431–445, 2020.",
    "O. Planchon and F. Darboux, “A fast, simple and versatile algorithm to fill the depressions of "
    "digital elevation models,” Catena, vol. 46, pp. 159–176, 2002.",
    "N. Trinh et al., “Flood risk in the Mekong and Chao Phraya deltas,” 2017.",
]


if __name__ == "__main__":
    build()
